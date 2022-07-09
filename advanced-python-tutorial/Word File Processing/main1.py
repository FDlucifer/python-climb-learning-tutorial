from docx.api import Document
from docx.shared import Pt

document = Document('test.docx')
all_text = ""
all_16pt_text = ""

for p in document.paragraphs:
    if p.style.name.startswith("Heading") or p.style.name == "Title":
        print(p.text)

for table in document.tables:
    print('NEW TABLE')
    for row in table.rows:
        print("|".join([cell.text for cell in row.cells]))

for p in document.paragraphs:
    all_text += p.text
    all_text += "\n"

print(all_text)

for p in document.paragraphs:
    for run in p.runs:
        if run.font.size == Pt(16):
            all_16pt_text += p.text
            all_16pt_text += "\n"

print(all_16pt_text)