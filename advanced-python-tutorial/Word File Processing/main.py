# pip install python-docx

from docx import Document

document = Document()
document.add_heading("fuck world", 0)
p = document.add_paragraph("this is a sample text!")
p.add_run("this text is bold.").bold = True
p.add_run("this text is italic.").italic = True

document.add_paragraph("this is item one", style="List Bullet")
document.add_paragraph("this is item two", style="List Bullet")
document.add_paragraph("this is item three", style="List Bullet")
document.add_paragraph("this is item four", style="List Bullet")
document.add_paragraph("this is item five", style="List Bullet")

table_header = ["Name", "Age", "Job"]

some_data = [
    ["John", 46, "Programmer"],
    ["Mary", 55, "Programmer"],
    ["Anna", 27, "Accountant"],
    ["Bob", 50, "Chef"],
]

table = document.add_table(rows=1, cols=3)
for i in range(3):
    table.rows[0].cells[i].text = table_header[i]

for name, age, job in some_data:
    cells = table.add_row().cells
    cells[0].text = name
    cells[1].text = str(age)
    cells[2].text = job

document.add_page_break()
document.add_paragraph("hello new page")
document.add_picture("Icon1024.png")
document.save("test.docx")